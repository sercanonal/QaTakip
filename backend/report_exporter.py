"""
Report Export Utilities
Generate professional presentation-quality reports in Word, Excel, and PDF formats
"""

import io
import logging
import os
from datetime import datetime, timedelta, timezone
from typing import List, Dict, Any
from collections import defaultdict
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, KeepTogether
from reportlab.platypus import Image as RLImage
from reportlab.graphics.shapes import Drawing, Rect, String, Circle
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.chart import PieChart, BarChart, Reference
from openpyxl.chart.label import DataLabelList
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)

# Register DejaVu fonts for Turkish character support
try:
    DEJAVU_PATH = '/usr/share/fonts/truetype/dejavu/'
    if os.path.exists(DEJAVU_PATH + 'DejaVuSans.ttf'):
        pdfmetrics.registerFont(TTFont('DejaVuSans', DEJAVU_PATH + 'DejaVuSans.ttf'))
        pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', DEJAVU_PATH + 'DejaVuSans-Bold.ttf'))
        TURKISH_FONT = 'DejaVuSans'
        TURKISH_FONT_BOLD = 'DejaVuSans-Bold'
        logger.info("DejaVu fonts registered for Turkish support")
    else:
        TURKISH_FONT = 'Helvetica'
        TURKISH_FONT_BOLD = 'Helvetica-Bold'
        logger.warning("DejaVu fonts not found, using Helvetica (Turkish chars may not render)")
except Exception as e:
    TURKISH_FONT = 'Helvetica'
    TURKISH_FONT_BOLD = 'Helvetica-Bold'
    logger.warning(f"Font registration failed: {e}")

# Intertech brand colors
BRAND_PRIMARY = '#E11D48'  # Rose/Red
BRAND_DARK = '#18181B'  # Dark background
BRAND_SECONDARY = '#27272A'  # Secondary dark
BRAND_SUCCESS = '#10B981'  # Green
BRAND_WARNING = '#F59E0B'  # Orange
BRAND_INFO = '#3B82F6'  # Blue

# Category colors
CATEGORY_COLORS = {
    'api-test': '#3B82F6',  # Blue
    'ui-test': '#10B981',  # Green
    'regression': '#F59E0B',  # Orange
    'bug-tracking': '#EF4444',  # Red
    'documentation': '#8B5CF6',  # Purple
}

# Status labels in Turkish
STATUS_LABELS = {
    'backlog': 'Backlog',
    'today_planned': 'Bugün Planlanan',
    'in_progress': 'Devam Ediyor',
    'blocked': 'Bloke',
    'completed': 'Tamamlandı'
}

# Priority labels in Turkish
PRIORITY_LABELS = {
    'low': 'Düşük',
    'medium': 'Orta',
    'high': 'Yüksek',
    'critical': 'Kritik'
}

# Category labels in Turkish
CATEGORY_LABELS = {
    'api-test': 'API Testi',
    'ui-test': 'UI Testi',
    'regression': 'Regresyon',
    'bug-tracking': 'Bug Tracking',
    'documentation': 'Test Dokümantasyonu'
}


class ReportExporter:
    """Generate professional presentation-quality reports in multiple formats"""
    
    @staticmethod
    def _calculate_monthly_stats(tasks: List[Dict]) -> Dict:
        """Calculate monthly task statistics"""
        monthly_data = defaultdict(lambda: {'created': 0, 'completed': 0})
        
        for task in tasks:
            created_at = task.get('created_at', '')
            completed_at = task.get('completed_at')
            
            if created_at:
                try:
                    created_date = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                    month_key = created_date.strftime('%Y-%m')
                    monthly_data[month_key]['created'] += 1
                except:
                    pass
            
            if completed_at:
                try:
                    completed_date = datetime.fromisoformat(completed_at.replace('Z', '+00:00'))
                    month_key = completed_date.strftime('%Y-%m')
                    monthly_data[month_key]['completed'] += 1
                except:
                    pass
        
        return dict(monthly_data)
    
    @staticmethod
    def _calculate_category_stats(tasks: List[Dict]) -> Dict:
        """Calculate category-based statistics"""
        category_data = defaultdict(lambda: {'total': 0, 'completed': 0, 'in_progress': 0})
        
        for task in tasks:
            category = task.get('category_id', 'other')
            status = task.get('status', '')
            
            category_data[category]['total'] += 1
            
            if status == 'completed':
                category_data[category]['completed'] += 1
            elif status == 'in_progress':
                category_data[category]['in_progress'] += 1
        
        return dict(category_data)
    
    @staticmethod
    def _calculate_priority_stats(tasks: List[Dict]) -> Dict:
        """Calculate priority-based statistics"""
        priority_data = defaultdict(int)
        
        for task in tasks:
            priority = task.get('priority', 'medium')
            priority_data[priority] += 1
        
        return dict(priority_data)
    
    @staticmethod
    def generate_pdf_report(data: Dict[str, Any]) -> bytes:
        """Generate professional canvas-style PDF report with modern design and Turkish support"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=A4,
            topMargin=0.5*cm,
            bottomMargin=0.5*cm,
            leftMargin=1*cm,
            rightMargin=1*cm
        )
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles - Modern dark theme with Turkish font support
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=32,
            textColor=colors.HexColor('#FFFFFF'),
            spaceAfter=8,
            alignment=1,
            fontName=TURKISH_FONT_BOLD
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=14,
            textColor=colors.HexColor('#A78BFA'),
            spaceAfter=20,
            alignment=1,
            fontName=TURKISH_FONT
        )
        
        section_title_style = ParagraphStyle(
            'SectionTitle',
            parent=styles['Heading2'],
            fontSize=18,
            textColor=colors.HexColor('#8B5CF6'),
            spaceBefore=25,
            spaceAfter=15,
            fontName=TURKISH_FONT_BOLD
        )
        
        kpi_label_style = ParagraphStyle(
            'KPILabel',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#71717A'),
            alignment=1,
            fontName=TURKISH_FONT
        )
        
        kpi_value_style = ParagraphStyle(
            'KPIValue',
            parent=styles['Normal'],
            fontSize=28,
            textColor=colors.HexColor('#FFFFFF'),
            alignment=1,
            fontName=TURKISH_FONT_BOLD
        )
        
        # Normal text style with Turkish support
        normal_turkish = ParagraphStyle(
            'NormalTurkish',
            parent=styles['Normal'],
            fontName=TURKISH_FONT,
            fontSize=10
        )
        
        # Get user info
        user_name = data.get('user_name', 'Kullanıcı')
        period_label = data.get('period_label', 'Son 30 Gün')
        stats = data.get('stats', {})
        tasks = data.get('tasks', [])
        
        # ===== HEADER BANNER =====
        # Create a dark gradient header background
        header_drawing = Drawing(525, 140)
        
        # Main dark background
        header_bg = Rect(0, 0, 525, 140, fillColor=colors.HexColor('#18181B'), strokeColor=None)
        header_drawing.add(header_bg)
        
        # Accent gradient overlay (purple)
        accent_rect = Rect(0, 0, 180, 140, fillColor=colors.HexColor('#4C1D95'), strokeColor=None)
        header_drawing.add(accent_rect)
        
        # Decorative circles
        circle1 = Circle(450, 100, 60, fillColor=colors.HexColor('#7C3AED'), fillOpacity=0.3, strokeColor=None)
        header_drawing.add(circle1)
        circle2 = Circle(480, 40, 40, fillColor=colors.HexColor('#A78BFA'), fillOpacity=0.2, strokeColor=None)
        header_drawing.add(circle2)
        
        story.append(header_drawing)
        
        # Title text over the header
        story.append(Spacer(1, -120))
        story.append(Paragraph("QA Hub", title_style))
        story.append(Paragraph("Performans Raporu", subtitle_style))
        
        # User info box
        info_style = ParagraphStyle(
            'InfoStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#D4D4D8'),
            alignment=1,
            fontName=TURKISH_FONT
        )
        story.append(Paragraph(f"<b>{user_name}</b> | {datetime.now().strftime('%d.%m.%Y')} | {period_label}", info_style))
        story.append(Spacer(1, 35))
        
        # ===== KPI CARDS SECTION =====
        if stats:
            story.append(Paragraph("Performans Gostergeleri", section_title_style))
            
            total = stats.get('total_tasks', 0)
            completed = stats.get('completed_tasks', 0)
            in_progress = stats.get('in_progress_tasks', 0)
            todo = stats.get('todo_tasks', 0)
            overdue = stats.get('overdue_tasks', 0)
            completion_rate = stats.get('completion_rate', 0)
            
            # Modern KPI Card design
            kpi_colors = [
                ('#8B5CF6', '#A78BFA'),  # Purple - Total
                ('#10B981', '#34D399'),  # Green - Completed
                ('#3B82F6', '#60A5FA'),  # Blue - In Progress
                ('#F59E0B', '#FBBF24'),  # Amber - Waiting
                ('#EF4444', '#F87171'),  # Red - Overdue
            ]
            
            # Create KPI cards as a styled table
            kpi_data = [
                [
                    Paragraph(f'<font size="24" color="#8B5CF6"><b>{total}</b></font>', normal_turkish),
                    Paragraph(f'<font size="24" color="#10B981"><b>{completed}</b></font>', normal_turkish),
                    Paragraph(f'<font size="24" color="#3B82F6"><b>{in_progress}</b></font>', normal_turkish),
                    Paragraph(f'<font size="24" color="#F59E0B"><b>{todo}</b></font>', normal_turkish),
                    Paragraph(f'<font size="24" color="#EF4444"><b>{overdue}</b></font>', normal_turkish),
                ],
                [
                    Paragraph('<font size="8" color="#71717A">TOPLAM</font>', normal_turkish),
                    Paragraph('<font size="8" color="#71717A">TAMAMLANAN</font>', normal_turkish),
                    Paragraph('<font size="8" color="#71717A">DEVAM EDEN</font>', normal_turkish),
                    Paragraph('<font size="8" color="#71717A">BEKLEYEN</font>', normal_turkish),
                    Paragraph('<font size="8" color="#71717A">GECIKMIS</font>', normal_turkish),
                ]
            ]
            
            kpi_table = Table(kpi_data, colWidths=[3.2*cm]*5)
            kpi_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#FAFAFA')),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 15),
                ('TOPPADDING', (0, 0), (-1, 0), 15),
                ('BOTTOMPADDING', (0, 1), (-1, 1), 10),
                ('TOPPADDING', (0, 1), (-1, 1), 5),
                ('ROUNDEDCORNERS', [8, 8, 8, 8]),
                ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#E5E7EB')),
            ]))
            story.append(kpi_table)
            story.append(Spacer(1, 0.8*cm))
            
            # Completion Rate Progress Bar - Modern Style
            story.append(Paragraph(f'<font size="11" color="#374151"><b>Tamamlanma Oranı: </b></font><font size="14" color="#8B5CF6"><b>%{completion_rate}</b></font>', styles['Normal']))
            story.append(Spacer(1, 0.3*cm))
            
            # Progress bar visual
            filled_width = min(completion_rate / 100 * 17, 17)
            empty_width = 17 - filled_width
            
            progress_drawing = Drawing(485, 20)
            # Background bar
            bg_bar = Rect(0, 5, 485, 10, fillColor=colors.HexColor('#E5E7EB'), strokeColor=None, rx=5, ry=5)
            progress_drawing.add(bg_bar)
            # Filled bar
            if filled_width > 0:
                filled_bar = Rect(0, 5, filled_width * 28.5, 10, fillColor=colors.HexColor('#8B5CF6'), strokeColor=None, rx=5, ry=5)
                progress_drawing.add(filled_bar)
            story.append(progress_drawing)
            story.append(Spacer(1, 1*cm))
        
        # ===== CATEGORY BREAKDOWN =====
        if tasks:
            story.append(Paragraph("Kategori Bazli Analiz", section_title_style))
            
            category_stats = ReportExporter._calculate_category_stats(tasks)
            
            if category_stats:
                cat_data = [
                    [
                        Paragraph('<font size="9" color="#FFFFFF"><b>Kategori</b></font>', normal_turkish),
                        Paragraph('<font size="9" color="#FFFFFF"><b>Toplam</b></font>', normal_turkish),
                        Paragraph('<font size="9" color="#FFFFFF"><b>Tamamlanan</b></font>', normal_turkish),
                        Paragraph('<font size="9" color="#FFFFFF"><b>Devam Eden</b></font>', normal_turkish),
                        Paragraph('<font size="9" color="#FFFFFF"><b>Basari %</b></font>', normal_turkish),
                    ]
                ]
                
                cat_colors = {
                    'api-test': '#3B82F6',
                    'ui-test': '#10B981',
                    'regression': '#F59E0B',
                    'bug-tracking': '#EF4444',
                    'documentation': '#8B5CF6',
                }
                
                for cat_id, cat_stats in category_stats.items():
                    cat_name = CATEGORY_LABELS.get(cat_id, cat_id)
                    cat_color = cat_colors.get(cat_id, '#6B7280')
                    total = cat_stats['total']
                    completed = cat_stats['completed']
                    in_prog = cat_stats['in_progress']
                    rate = round(completed / total * 100, 1) if total > 0 else 0
                    
                    # Color-coded rate
                    rate_color = '#10B981' if rate >= 70 else '#F59E0B' if rate >= 40 else '#EF4444'
                    
                    cat_data.append([
                        Paragraph(f'<font size="9" color="{cat_color}"><b>* </b></font><font size="9">{cat_name}</font>', normal_turkish),
                        Paragraph(f'<font size="10"><b>{total}</b></font>', normal_turkish),
                        Paragraph(f'<font size="10" color="#10B981"><b>{completed}</b></font>', normal_turkish),
                        Paragraph(f'<font size="10" color="#3B82F6"><b>{in_prog}</b></font>', normal_turkish),
                        Paragraph(f'<font size="10" color="{rate_color}"><b>%{rate}</b></font>', normal_turkish),
                    ])
                
                cat_table = Table(cat_data, colWidths=[5.5*cm, 2.5*cm, 2.8*cm, 2.8*cm, 2.5*cm])
                cat_table.setStyle(TableStyle([
                    # Header
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7C3AED')),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('TOPPADDING', (0, 0), (-1, 0), 12),
                    # Data rows
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 10),
                    ('TOPPADDING', (0, 1), (-1, -1), 10),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F9FAFB')]),
                    # Rounded corners effect with box
                    ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#E5E7EB')),
                    ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#7C3AED')),
                ]))
                story.append(cat_table)
                story.append(Spacer(1, 1*cm))
        
        # ===== PRIORITY BREAKDOWN =====
        if tasks:
            story.append(Paragraph("Oncelik Dagilimi", section_title_style))
            
            priority_stats = ReportExporter._calculate_priority_stats(tasks)
            
            if priority_stats:
                priority_colors = {
                    'critical': '#DC2626',
                    'high': '#EA580C',
                    'medium': '#CA8A04',
                    'low': '#65A30D'
                }
                
                priority_markers = {
                    'critical': '[!]',
                    'high': '[H]',
                    'medium': '[M]',
                    'low': '[L]'
                }
                
                priority_data = [
                    [
                        Paragraph('<font size="9" color="#FFFFFF"><b>Oncelik</b></font>', normal_turkish),
                        Paragraph('<font size="9" color="#FFFFFF"><b>Gorev Sayisi</b></font>', normal_turkish),
                        Paragraph('<font size="9" color="#FFFFFF"><b>Oran</b></font>', normal_turkish),
                        Paragraph('<font size="9" color="#FFFFFF"><b>Gorsel</b></font>', normal_turkish),
                    ]
                ]
                total_tasks = sum(priority_stats.values())
                
                for priority in ['critical', 'high', 'medium', 'low']:
                    if priority in priority_stats:
                        count = priority_stats[priority]
                        rate = round(count / total_tasks * 100, 1) if total_tasks > 0 else 0
                        color = priority_colors.get(priority, '#6B7280')
                        marker = priority_markers.get(priority, '*')
                        
                        # Create visual bar representation
                        bar_length = int(rate / 5)  # Scale to fit
                        bar = '|' * bar_length + '.' * (20 - bar_length)
                        
                        priority_data.append([
                            Paragraph(f'<font size="10" color="{color}"><b>{marker} {PRIORITY_LABELS.get(priority, priority)}</b></font>', normal_turkish),
                            Paragraph(f'<font size="12"><b>{count}</b></font>', normal_turkish),
                            Paragraph(f'<font size="10" color="{color}"><b>%{rate}</b></font>', normal_turkish),
                            Paragraph(f'<font size="8" color="{color}">{bar}</font>', normal_turkish),
                        ])
                
                if len(priority_data) > 1:
                    priority_table = Table(priority_data, colWidths=[4*cm, 3*cm, 2.5*cm, 6.5*cm])
                    priority_table.setStyle(TableStyle([
                        # Header
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#18181B')),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('ALIGN', (3, 1), (3, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('TOPPADDING', (0, 0), (-1, 0), 12),
                        # Data rows
                        ('BOTTOMPADDING', (0, 1), (-1, -1), 10),
                        ('TOPPADDING', (0, 1), (-1, -1), 10),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#FAFAFA')]),
                        # Box
                        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#E5E7EB')),
                        ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#18181B')),
                    ]))
                    story.append(priority_table)
                    story.append(Spacer(1, 1*cm))
        
        # ===== MONTHLY PERFORMANCE =====
        if tasks:
            monthly_stats = ReportExporter._calculate_monthly_stats(tasks)
            
            if monthly_stats:
                story.append(Paragraph("Aylik Performans Trendi", section_title_style))
                
                monthly_data = [
                    [
                        Paragraph('<font size="9" color="#FFFFFF"><b>Ay</b></font>', normal_turkish),
                        Paragraph('<font size="9" color="#FFFFFF"><b>Olusturulan</b></font>', normal_turkish),
                        Paragraph('<font size="9" color="#FFFFFF"><b>Tamamlanan</b></font>', normal_turkish),
                        Paragraph('<font size="9" color="#FFFFFF"><b>Performans</b></font>', normal_turkish),
                        Paragraph('<font size="9" color="#FFFFFF"><b>Trend</b></font>', normal_turkish),
                    ]
                ]
                
                sorted_months = sorted(monthly_stats.keys(), reverse=True)[:6]  # Son 6 ay
                prev_perf = None
                
                for month in sorted_months:
                    stats_m = monthly_stats[month]
                    created = stats_m['created']
                    completed = stats_m['completed']
                    perf = round(completed / created * 100, 1) if created > 0 else 0
                    
                    # Trend indicator
                    if prev_perf is not None:
                        if perf > prev_perf:
                            trend = '++ UP'
                            trend_color = '#10B981'
                        elif perf < prev_perf:
                            trend = '-- DN'
                            trend_color = '#EF4444'
                        else:
                            trend = '➡️ ='
                            trend_color = '#6B7280'
                    else:
                        trend = '—'
                        trend_color = '#6B7280'
                    prev_perf = perf
                    
                    # Format month name in Turkish
                    try:
                        month_date = datetime.strptime(month, '%Y-%m')
                        month_names = ['Ocak', 'Şubat', 'Mart', 'Nisan', 'Mayıs', 'Haziran',
                                      'Temmuz', 'Ağustos', 'Eylül', 'Ekim', 'Kasım', 'Aralık']
                        month_name = f"{month_names[month_date.month - 1]} {month_date.year}"
                    except:
                        month_name = month
                    
                    # Performance color
                    perf_color = '#10B981' if perf >= 70 else '#F59E0B' if perf >= 40 else '#EF4444'
                    
                    monthly_data.append([
                        Paragraph(f'<font size="9"><b>{month_name}</b></font>', styles['Normal']),
                        Paragraph(f'<font size="10" color="#3B82F6"><b>{created}</b></font>', styles['Normal']),
                        Paragraph(f'<font size="10" color="#10B981"><b>{completed}</b></font>', styles['Normal']),
                        Paragraph(f'<font size="10" color="{perf_color}"><b>%{perf}</b></font>', styles['Normal']),
                        Paragraph(f'<font size="9" color="{trend_color}">{trend}</font>', styles['Normal']),
                    ])
                
                if len(monthly_data) > 1:
                    monthly_table = Table(monthly_data, colWidths=[4.5*cm, 3*cm, 3*cm, 3*cm, 2.5*cm])
                    monthly_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3B82F6')),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('TOPPADDING', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 1), (-1, -1), 10),
                        ('TOPPADDING', (0, 1), (-1, -1), 10),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#EFF6FF')]),
                        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#E5E7EB')),
                        ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#3B82F6')),
                    ]))
                    story.append(monthly_table)
                    story.append(Spacer(1, 1*cm))
        
        # ===== TASK LIST =====
        if tasks:
            story.append(PageBreak())
            
            # Page header for tasks
            task_header_drawing = Drawing(525, 60)
            task_header_bg = Rect(0, 0, 525, 60, fillColor=colors.HexColor('#F3F4F6'), strokeColor=None)
            task_header_drawing.add(task_header_bg)
            accent_line = Rect(0, 0, 6, 60, fillColor=colors.HexColor('#8B5CF6'), strokeColor=None)
            task_header_drawing.add(accent_line)
            story.append(task_header_drawing)
            
            story.append(Spacer(1, -45))
            story.append(Paragraph("📋 Görev Detayları", section_title_style))
            story.append(Spacer(1, 10))
            
            task_data = [
                [
                    Paragraph('<font size="8" color="#FFFFFF"><b>#</b></font>', styles['Normal']),
                    Paragraph('<font size="8" color="#FFFFFF"><b>Başlık</b></font>', styles['Normal']),
                    Paragraph('<font size="8" color="#FFFFFF"><b>Kategori</b></font>', styles['Normal']),
                    Paragraph('<font size="8" color="#FFFFFF"><b>Durum</b></font>', styles['Normal']),
                    Paragraph('<font size="8" color="#FFFFFF"><b>Öncelik</b></font>', styles['Normal']),
                ]
            ]
            
            status_colors = {
                'completed': '#10B981',
                'in_progress': '#3B82F6',
                'blocked': '#EF4444',
                'backlog': '#6B7280',
                'today_planned': '#F59E0B'
            }
            
            priority_colors = {
                'critical': '#DC2626',
                'high': '#EA580C',
                'medium': '#CA8A04',
                'low': '#65A30D'
            }
            
            for idx, task in enumerate(tasks[:30], 1):  # Limit to 30 for PDF
                title = task.get('title', '')[:40]
                if len(task.get('title', '')) > 40:
                    title += '...'
                
                category = CATEGORY_LABELS.get(task.get('category_id', ''), task.get('category_id', ''))
                status = task.get('status', '')
                status_label = STATUS_LABELS.get(status, status)
                status_color = status_colors.get(status, '#6B7280')
                priority = task.get('priority', '')
                priority_label = PRIORITY_LABELS.get(priority, priority)
                priority_color = priority_colors.get(priority, '#6B7280')
                
                task_data.append([
                    Paragraph(f'<font size="8">{idx}</font>', styles['Normal']),
                    Paragraph(f'<font size="8">{title}</font>', styles['Normal']),
                    Paragraph(f'<font size="8">{category}</font>', styles['Normal']),
                    Paragraph(f'<font size="8" color="{status_color}"><b>{status_label}</b></font>', styles['Normal']),
                    Paragraph(f'<font size="8" color="{priority_color}"><b>{priority_label}</b></font>', styles['Normal']),
                ])
            
            task_table = Table(task_data, colWidths=[1*cm, 8*cm, 3*cm, 2.5*cm, 2*cm])
            task_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#18181B')),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('TOPPADDING', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F9FAFB')]),
                ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#E5E7EB')),
                ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#8B5CF6')),
            ]))
            story.append(task_table)
        
        # ===== FOOTER =====
        story.append(Spacer(1, 1.5*cm))
        
        footer_drawing = Drawing(525, 50)
        footer_bg = Rect(0, 0, 525, 50, fillColor=colors.HexColor('#18181B'), strokeColor=None)
        footer_drawing.add(footer_bg)
        footer_text = String(262, 30, "QA Hub - Performans Raporu", fontSize=10, fillColor=colors.HexColor('#A78BFA'), textAnchor='middle')
        footer_drawing.add(footer_text)
        footer_date = String(262, 12, f"Oluşturulma: {datetime.now().strftime('%d.%m.%Y %H:%M')} | Intertech QA Ekibi", fontSize=8, fillColor=colors.HexColor('#71717A'), textAnchor='middle')
        footer_drawing.add(footer_date)
        story.append(footer_drawing)
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    @staticmethod
    def generate_excel_report(data: Dict[str, Any]) -> bytes:
        """Generate professional canvas-style Excel report with charts"""
        wb = Workbook()
        
        # Modern Styles
        title_font = Font(size=28, bold=True, color="8B5CF6")
        subtitle_font = Font(size=12, color="71717A")
        header_font = Font(size=11, bold=True, color="FFFFFF")
        value_font = Font(size=14, bold=True, color="18181B")
        
        # Modern color fills
        header_fill = PatternFill(start_color="18181B", end_color="18181B", fill_type="solid")
        purple_fill = PatternFill(start_color="7C3AED", end_color="7C3AED", fill_type="solid")
        blue_fill = PatternFill(start_color="3B82F6", end_color="3B82F6", fill_type="solid")
        green_fill = PatternFill(start_color="10B981", end_color="10B981", fill_type="solid")
        amber_fill = PatternFill(start_color="F59E0B", end_color="F59E0B", fill_type="solid")
        red_fill = PatternFill(start_color="EF4444", end_color="EF4444", fill_type="solid")
        light_fill = PatternFill(start_color="F9FAFB", end_color="F9FAFB", fill_type="solid")
        purple_light = PatternFill(start_color="F3E8FF", end_color="F3E8FF", fill_type="solid")
        
        thin_border = Border(
            left=Side(style='thin', color='E5E7EB'),
            right=Side(style='thin', color='E5E7EB'),
            top=Side(style='thin', color='E5E7EB'),
            bottom=Side(style='thin', color='E5E7EB')
        )
        
        user_name = data.get('user_name', 'Kullanıcı')
        period_label = data.get('period_label', 'Son 30 Gün')
        stats = data.get('stats', {})
        tasks = data.get('tasks', [])
        
        # ===== DASHBOARD SHEET =====
        ws_dash = wb.active
        ws_dash.title = "📊 Dashboard"
        
        # Header section with gradient effect
        ws_dash.merge_cells('A1:H3')
        ws_dash['A1'] = "🚀 QA Hub - Performans Raporu"
        ws_dash['A1'].font = title_font
        ws_dash['A1'].alignment = Alignment(horizontal='center', vertical='center')
        ws_dash['A1'].fill = PatternFill(start_color="18181B", end_color="18181B", fill_type="solid")
        
        # Apply dark header to entire row range
        for col in range(1, 9):
            for row in range(1, 4):
                ws_dash.cell(row=row, column=col).fill = PatternFill(start_color="18181B", end_color="18181B", fill_type="solid")
        
        ws_dash.merge_cells('A4:H4')
        ws_dash['A4'] = f"👤 {user_name}  |  📅 {datetime.now().strftime('%d.%m.%Y')}  |  📆 {period_label}"
        ws_dash['A4'].alignment = Alignment(horizontal='center')
        ws_dash['A4'].font = subtitle_font
        ws_dash['A4'].fill = purple_light
        
        # Row heights
        ws_dash.row_dimensions[1].height = 25
        ws_dash.row_dimensions[2].height = 25
        ws_dash.row_dimensions[3].height = 25
        ws_dash.row_dimensions[4].height = 30
        
        # KPI Section
        ws_dash.merge_cells('A6:H6')
        ws_dash['A6'] = "📈 PERFORMANS GÖSTERGELERİ"
        ws_dash['A6'].font = Font(size=14, bold=True, color="8B5CF6")
        ws_dash['A6'].fill = light_fill
        
        if stats:
            kpi_data = [
                ('🎯 Toplam', stats.get('total_tasks', 0), purple_fill),
                ('✅ Tamamlanan', stats.get('completed_tasks', 0), green_fill),
                ('🔄 Devam Eden', stats.get('in_progress_tasks', 0), blue_fill),
                ('⏳ Bekleyen', stats.get('todo_tasks', 0), amber_fill),
                ('⚠️ Gecikmiş', stats.get('overdue_tasks', 0), red_fill),
                ('📊 Başarı %', f"%{stats.get('completion_rate', 0)}", purple_fill),
            ]
            
            for col, (label, value, fill) in enumerate(kpi_data, 1):
                # Header cell
                header_cell = ws_dash.cell(row=8, column=col, value=label)
                header_cell.font = Font(size=10, bold=True, color="FFFFFF")
                header_cell.fill = fill
                header_cell.alignment = Alignment(horizontal='center')
                header_cell.border = thin_border
                
                # Value cell
                value_cell = ws_dash.cell(row=9, column=col, value=value)
                value_cell.font = Font(size=20, bold=True)
                value_cell.alignment = Alignment(horizontal='center')
                value_cell.border = thin_border
                
                ws_dash.column_dimensions[get_column_letter(col)].width = 14
            
            ws_dash.row_dimensions[8].height = 30
            ws_dash.row_dimensions[9].height = 40
        
        # Category Analysis Section
        if tasks:
            ws_dash.merge_cells('A12:H12')
            ws_dash['A12'] = "🎯 KATEGORİ ANALİZİ"
            ws_dash['A12'].font = Font(size=14, bold=True, color="7C3AED")
            ws_dash['A12'].fill = purple_light
            
            category_stats = ReportExporter._calculate_category_stats(tasks)
            
            cat_headers = ['Kategori', 'Toplam', 'Tamamlanan', 'Devam Eden', 'Başarı %']
            cat_header_fills = [purple_fill, header_fill, green_fill, blue_fill, purple_fill]
            
            for col, (header, fill) in enumerate(zip(cat_headers, cat_header_fills), 1):
                cell = ws_dash.cell(row=14, column=col, value=header)
                cell.font = header_font
                cell.fill = fill
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
            
            row = 15
            cat_icons = {
                'api-test': '🔌',
                'ui-test': '🖥️',
                'regression': '🔄',
                'bug-tracking': '🐛',
                'documentation': '📝',
            }
            
            for cat_id, cat_stats in category_stats.items():
                icon = cat_icons.get(cat_id, '📋')
                cat_name = CATEGORY_LABELS.get(cat_id, cat_id)
                total = cat_stats['total']
                completed = cat_stats['completed']
                in_prog = cat_stats['in_progress']
                rate = round(completed / total * 100, 1) if total > 0 else 0
                
                ws_dash.cell(row=row, column=1, value=f"{icon} {cat_name}").border = thin_border
                ws_dash.cell(row=row, column=2, value=total).border = thin_border
                ws_dash.cell(row=row, column=3, value=completed).border = thin_border
                ws_dash.cell(row=row, column=4, value=in_prog).border = thin_border
                ws_dash.cell(row=row, column=5, value=f"%{rate}").border = thin_border
                
                # Color code completion rate
                rate_cell = ws_dash.cell(row=row, column=5)
                if rate >= 70:
                    rate_cell.font = Font(bold=True, color="10B981")
                elif rate >= 40:
                    rate_cell.font = Font(bold=True, color="F59E0B")
                else:
                    rate_cell.font = Font(bold=True, color="EF4444")
                
                for col in range(1, 6):
                    ws_dash.cell(row=row, column=col).alignment = Alignment(horizontal='center')
                    if row % 2 == 0:
                        ws_dash.cell(row=row, column=col).fill = light_fill
                
                row += 1
            
            # Add pie chart
            if len(category_stats) > 0:
                pie = PieChart()
                pie.title = "Kategori Dağılımı"
                labels = Reference(ws_dash, min_col=1, min_row=15, max_row=row-1)
                data_ref = Reference(ws_dash, min_col=2, min_row=14, max_row=row-1)
                pie.add_data(data_ref, titles_from_data=True)
                pie.set_categories(labels)
                pie.width = 14
                pie.height = 10
                ws_dash.add_chart(pie, "G12")
        
        # Column widths for dashboard
        ws_dash.column_dimensions['A'].width = 20
        for col in range(2, 9):
            ws_dash.column_dimensions[get_column_letter(col)].width = 14
        
        # ===== TASKS SHEET =====
        ws_tasks = wb.create_sheet("📋 Görevler")
        
        # Header styling for tasks sheet
        task_headers = ["#", "📝 Başlık", "📁 Kategori", "📊 Durum", "⚡ Öncelik", "📅 Oluşturulma", "✅ Tamamlanma"]
        header_fills_task = [header_fill, purple_fill, blue_fill, green_fill, amber_fill, header_fill, green_fill]
        
        for col, (header, fill) in enumerate(zip(task_headers, header_fills_task), 1):
            cell = ws_tasks.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        
        ws_tasks.row_dimensions[1].height = 30
        
        status_colors = {
            'completed': '10B981',
            'in_progress': '3B82F6',
            'blocked': 'EF4444',
            'backlog': '6B7280',
            'today_planned': 'F59E0B'
        }
        
        priority_colors = {
            'critical': 'DC2626',
            'high': 'EA580C',
            'medium': 'CA8A04',
            'low': '65A30D'
        }
        
        for row_num, task in enumerate(tasks, 2):
            ws_tasks.cell(row=row_num, column=1, value=row_num-1).border = thin_border
            ws_tasks.cell(row=row_num, column=2, value=task.get('title', '')).border = thin_border
            ws_tasks.cell(row=row_num, column=3, value=CATEGORY_LABELS.get(task.get('category_id', ''), task.get('category_id', ''))).border = thin_border
            
            status = task.get('status', '')
            status_cell = ws_tasks.cell(row=row_num, column=4, value=STATUS_LABELS.get(status, status))
            status_cell.border = thin_border
            status_cell.font = Font(bold=True, color=status_colors.get(status, '6B7280'))
            
            priority = task.get('priority', '')
            priority_cell = ws_tasks.cell(row=row_num, column=5, value=PRIORITY_LABELS.get(priority, priority))
            priority_cell.border = thin_border
            priority_cell.font = Font(bold=True, color=priority_colors.get(priority, '6B7280'))
            
            created = task.get('created_at', '')
            if created:
                try:
                    created_dt = datetime.fromisoformat(created.replace('Z', '+00:00'))
                    created = created_dt.strftime('%d.%m.%Y')
                except:
                    pass
            ws_tasks.cell(row=row_num, column=6, value=created).border = thin_border
            
            completed = task.get('completed_at', '')
            if completed:
                try:
                    completed_dt = datetime.fromisoformat(completed.replace('Z', '+00:00'))
                    completed = completed_dt.strftime('%d.%m.%Y')
                except:
                    pass
            ws_tasks.cell(row=row_num, column=7, value=completed or '-').border = thin_border
            
            for col in range(1, 8):
                ws_tasks.cell(row=row_num, column=col).alignment = Alignment(horizontal='center')
                if row_num % 2 == 0:
                    ws_tasks.cell(row=row_num, column=col).fill = light_fill
        
        # Column widths for tasks
        ws_tasks.column_dimensions['A'].width = 5
        ws_tasks.column_dimensions['B'].width = 45
        ws_tasks.column_dimensions['C'].width = 18
        ws_tasks.column_dimensions['D'].width = 15
        ws_tasks.column_dimensions['E'].width = 12
        ws_tasks.column_dimensions['F'].width = 14
        ws_tasks.column_dimensions['G'].width = 14
        
        # ===== MONTHLY STATS SHEET =====
        ws_monthly = wb.create_sheet("📈 Aylık Trend")
        
        monthly_stats = ReportExporter._calculate_monthly_stats(tasks)
        
        monthly_headers = ['📅 Ay', '📊 Oluşturulan', '✅ Tamamlanan', '📈 Performans %']
        monthly_fills = [blue_fill, purple_fill, green_fill, blue_fill]
        
        for col, (header, fill) in enumerate(zip(monthly_headers, monthly_fills), 1):
            cell = ws_monthly.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        
        ws_monthly.row_dimensions[1].height = 30
        
        row = 2
        for month in sorted(monthly_stats.keys(), reverse=True):
            stats_m = monthly_stats[month]
            created = stats_m['created']
            completed = stats_m['completed']
            perf = round(completed / created * 100, 1) if created > 0 else 0
            
            try:
                month_date = datetime.strptime(month, '%Y-%m')
                month_names = ['Ocak', 'Şubat', 'Mart', 'Nisan', 'Mayıs', 'Haziran',
                              'Temmuz', 'Ağustos', 'Eylül', 'Ekim', 'Kasım', 'Aralık']
                month_name = f"{month_names[month_date.month - 1]} {month_date.year}"
            except:
                month_name = month
            
            ws_monthly.cell(row=row, column=1, value=month_name).border = thin_border
            ws_monthly.cell(row=row, column=2, value=created).border = thin_border
            ws_monthly.cell(row=row, column=3, value=completed).border = thin_border
            
            perf_cell = ws_monthly.cell(row=row, column=4, value=f"%{perf}")
            perf_cell.border = thin_border
            if perf >= 70:
                perf_cell.font = Font(bold=True, color="10B981")
            elif perf >= 40:
                perf_cell.font = Font(bold=True, color="F59E0B")
            else:
                perf_cell.font = Font(bold=True, color="EF4444")
            
            for col in range(1, 5):
                ws_monthly.cell(row=row, column=col).alignment = Alignment(horizontal='center')
                if row % 2 == 0:
                    ws_monthly.cell(row=row, column=col).fill = light_fill
            
            row += 1
        
        # Column widths for monthly
        for col in range(1, 5):
            ws_monthly.column_dimensions[get_column_letter(col)].width = 20
        
        # Add bar chart for monthly performance
        if row > 2:
            chart = BarChart()
            chart.title = "📈 Aylık Performans Trendi"
            chart.x_axis.title = "Ay"
            chart.y_axis.title = "Görev Sayısı"
            chart.style = 12
            
            data_ref = Reference(ws_monthly, min_col=2, max_col=3, min_row=1, max_row=row-1)
            cats = Reference(ws_monthly, min_col=1, min_row=2, max_row=row-1)
            
            chart.add_data(data_ref, titles_from_data=True)
            chart.set_categories(cats)
            chart.shape = 4
            chart.width = 18
            chart.height = 12
            
            ws_monthly.add_chart(chart, "F1")
        
        # Save
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    @staticmethod
    def generate_word_report(data: Dict[str, Any]) -> bytes:
        """Generate professional Word document report"""
        doc = Document()
        
        user_name = data.get('user_name', 'Kullanıcı')
        stats = data.get('stats', {})
        tasks = data.get('tasks', [])
        
        # ===== TITLE PAGE =====
        title = doc.add_heading('QA Task Manager', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(225, 29, 72)  # Brand primary
        
        subtitle = doc.add_paragraph('Performans Raporu')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(113, 113, 122)
        
        doc.add_paragraph()
        
        # Report info
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"Hazırlayan: {user_name}\n").bold = True
        info_para.add_run(f"Rapor Tarihi: {datetime.now().strftime('%d %B %Y')}\n")
        info_para.add_run("Dönem: Son 30 Gün")
        
        doc.add_page_break()
        
        # ===== EXECUTIVE SUMMARY =====
        summary_heading = doc.add_heading('Yönetici Özeti', level=1)
        for run in summary_heading.runs:
            run.font.color.rgb = RGBColor(225, 29, 72)
        
        if stats:
            # KPI Table
            kpi_table = doc.add_table(rows=2, cols=5)
            kpi_table.style = 'Table Grid'
            kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            headers = ['Toplam Görev', 'Tamamlanan', 'Devam Eden', 'Bekleyen', 'Gecikmiş']
            values = [
                str(stats.get('total_tasks', 0)),
                str(stats.get('completed_tasks', 0)),
                str(stats.get('in_progress_tasks', 0)),
                str(stats.get('todo_tasks', 0)),
                str(stats.get('overdue_tasks', 0))
            ]
            
            # Header row
            header_row = kpi_table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                # Set background color
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="18181B"/>')
                cell._tc.get_or_add_tcPr().append(shading)
            
            # Value row
            value_row = kpi_table.rows[1]
            for i, value in enumerate(values):
                cell = value_row.cells[i]
                cell.text = value
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(20)
            
            doc.add_paragraph()
            
            # Completion rate
            rate = stats.get('completion_rate', 0)
            rate_para = doc.add_paragraph()
            rate_para.add_run(f"Tamamlanma Oranı: ").bold = True
            rate_run = rate_para.add_run(f"%{rate}")
            rate_run.font.size = Pt(24)
            rate_run.font.bold = True
            rate_run.font.color.rgb = RGBColor(16, 185, 129)  # Green
        
        doc.add_paragraph()
        
        # ===== CATEGORY ANALYSIS =====
        if tasks:
            cat_heading = doc.add_heading('Kategori Bazlı Analiz', level=1)
            for run in cat_heading.runs:
                run.font.color.rgb = RGBColor(225, 29, 72)
            
            category_stats = ReportExporter._calculate_category_stats(tasks)
            
            if category_stats:
                cat_table = doc.add_table(rows=len(category_stats) + 1, cols=5)
                cat_table.style = 'Table Grid'
                
                # Header
                headers = ['Kategori', 'Toplam', 'Tamamlanan', 'Devam Eden', 'Tamamlanma %']
                header_row = cat_table.rows[0]
                for i, header in enumerate(headers):
                    cell = header_row.cells[i]
                    cell.text = header
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E11D48"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                
                # Data rows
                for row_idx, (cat_id, cat_stats) in enumerate(category_stats.items(), 1):
                    row = cat_table.rows[row_idx]
                    cat_name = CATEGORY_LABELS.get(cat_id, cat_id)
                    total = cat_stats['total']
                    completed = cat_stats['completed']
                    in_prog = cat_stats['in_progress']
                    rate = round(completed / total * 100, 1) if total > 0 else 0
                    
                    values = [cat_name, str(total), str(completed), str(in_prog), f'%{rate}']
                    for i, value in enumerate(values):
                        cell = row.cells[i]
                        cell.text = value
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # ===== MONTHLY PERFORMANCE =====
        if tasks:
            monthly_stats = ReportExporter._calculate_monthly_stats(tasks)
            
            if monthly_stats:
                monthly_heading = doc.add_heading('Aylık Performans', level=1)
                for run in monthly_heading.runs:
                    run.font.color.rgb = RGBColor(225, 29, 72)
                
                sorted_months = sorted(monthly_stats.keys(), reverse=True)[:6]
                
                monthly_table = doc.add_table(rows=len(sorted_months) + 1, cols=4)
                monthly_table.style = 'Table Grid'
                
                # Header
                headers = ['Ay', 'Oluşturulan', 'Tamamlanan', 'Performans']
                header_row = monthly_table.rows[0]
                for i, header in enumerate(headers):
                    cell = header_row.cells[i]
                    cell.text = header
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="3B82F6"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                
                # Data rows
                for row_idx, month in enumerate(sorted_months, 1):
                    row = monthly_table.rows[row_idx]
                    stats_m = monthly_stats[month]
                    created = stats_m['created']
                    completed = stats_m['completed']
                    perf = round(completed / created * 100, 1) if created > 0 else 0
                    
                    try:
                        month_date = datetime.strptime(month, '%Y-%m')
                        month_names = ['Ocak', 'Şubat', 'Mart', 'Nisan', 'Mayıs', 'Haziran',
                                      'Temmuz', 'Ağustos', 'Eylül', 'Ekim', 'Kasım', 'Aralık']
                        month_name = f"{month_names[month_date.month - 1]} {month_date.year}"
                    except:
                        month_name = month
                    
                    values = [month_name, str(created), str(completed), f'%{perf}']
                    for i, value in enumerate(values):
                        cell = row.cells[i]
                        cell.text = value
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # ===== TASK DETAILS =====
        if tasks:
            task_heading = doc.add_heading('Görev Detayları', level=1)
            for run in task_heading.runs:
                run.font.color.rgb = RGBColor(225, 29, 72)
            
            task_table = doc.add_table(rows=min(len(tasks), 20) + 1, cols=4)
            task_table.style = 'Table Grid'
            
            # Header
            headers = ['Başlık', 'Kategori', 'Durum', 'Öncelik']
            header_row = task_table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="18181B"/>')
                cell._tc.get_or_add_tcPr().append(shading)
            
            # Task rows
            for row_idx, task in enumerate(tasks[:20], 1):
                row = task_table.rows[row_idx]
                
                title = task.get('title', '')[:40]
                if len(task.get('title', '')) > 40:
                    title += '...'
                
                values = [
                    title,
                    CATEGORY_LABELS.get(task.get('category_id', ''), task.get('category_id', '')),
                    STATUS_LABELS.get(task.get('status', ''), task.get('status', '')),
                    PRIORITY_LABELS.get(task.get('priority', ''), task.get('priority', ''))
                ]
                
                for i, value in enumerate(values):
                    cell = row.cells[i]
                    cell.text = value
                    if i > 0:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ===== FOOTER =====
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("Bu rapor QA Task Manager tarafından otomatik olarak oluşturulmuştur.\nIntertech QA Ekibi")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(113, 113, 122)
        
        # Save
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


# Global exporter instance
report_exporter = ReportExporter()
